from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests
import replicate

# Loading the model
def generate_llama2_response(prompt_input):
    string_dialogue = "You are a helpful assistant. You do not respond as 'User' or pretend to be 'User'. You only respond once as 'Assistant'."
    for dict_message in st.session_state.messages:
        if dict_message["role"] == "user":
            string_dialogue += "User: " + dict_message["content"] + "\n\n"
        else:
            string_dialogue += "Assistant: " + dict_message["content"] + "\n\n"
    output = replicate.run('a16z-infra/llama13b-v2-chat:df7690f1994d94e96ad9d568eac121aecf50684a0b0963b25a41cc40061269e5', 
                           input={"prompt": f"{string_dialogue} {prompt_input} Assistant: ",
                                  "temperature":temperature, "top_p":top_p, "max_length":max_length, "repetition_penalty":1})
    return output
    
def get_src_original_url(query):
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': st.secrets["PEXELS_API_KEY"],
    }
    params = {
        'query': query,
        'per_page': 1,
    }
    response = requests.get(url, headers=headers, params=params)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            return photos[0]['src']['original']
    return None

def save_image_from_url(url, filename="temp_image.jpg"):
    response = requests.get(url, stream=True)
    response.raise_for_status()
    with open(filename, 'wb') as file:
        for chunk in response.iter_content(8192):
            file.write(chunk)

def create_word_docx(user_input, paragraph, image_filename):
    doc = Document()
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    doc.add_picture(image_filename, width=Inches(4))
    return doc

def main():
    st.title("Article Generator App using Llama 2")
    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")

    if len(user_input) > 0 and len(image_input) > 0:
        col1, col2, col3 = st.columns([1,2,1])
        with col1:
            st.subheader("Generated Content by Llama 2")
            prompt_template = f"You are a digital marketing and SEO expert and your task is to write an article on the given topic: {user_input}. The article must be under 800 words."
            
            result = ""  # Initialize the result variable
            try:
                llm_chain = load_llm(max_tokens=800, prompt_template=prompt_template)
                generator_output = llm_chain(user_input)
                result = next(generator_output, {}).get('content', "")
                
                if result:
                    st.info("Your article has been been generated successfully!")
                    st.write(result)
                else:
                    st.error("Your article couldn't be generated or is empty!")

            except Exception as e:
                st.error(f"An error occurred while generating the article: {e}")

        with col2:
            st.subheader("Fetched Image")
            image_url = get_src_original_url(image_input)
            if image_url:
                save_image_from_url(image_url)
                st.image(image_url)
            else:
                st.error("Couldn't fetch the image for the given topic.")

        with col3:
            st.subheader("Final Article to Download")
            if result and image_url:
                image_filename = "temp_image.jpg"
                doc = create_word_docx(user_input, result, image_filename)
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                st.download_button(
                    label='Download Word Document',
                    data=doc_buffer,
                    file_name='document.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
            else:
                st.error("Couldn't generate the Word document. Ensure both the article and image are generated successfully.")

if __name__ == "__main__":
    main()
