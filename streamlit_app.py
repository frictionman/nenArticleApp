
# Developed by Nen AI

import requests
from PIL import Image
import io
from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
from docx import Document
from docx.shared import Inches

# Fetching an image from Pexels
def fetch_photo(query: str):
    API_KEY = 'YOUR_PEXELS_API_KEY'  # Please replace with your Pexels API key
    URL = f"https://api.pexels.com/v1/search?query={query}&per_page=1"
    headers = {
        "Authorization": API_KEY
    }
    response = requests.get(URL, headers=headers)
    if response.status_code != 200:
        st.write("Error fetching image from Pexels. Please check the API key and query.")
        return None
    data = response.json()
    if not data["photos"]:
        st.write("No photos found for the given query on Pexels.")
        return None
    photo_url = data["photos"][0]["src"]["large"]
    return photo_url

# Loading the LLM model
def load_llm_model(api_key: str, model_size: str):
    assert model_size in ["Llama2-7B", "Llama2-13B"], "Invalid model size provided."
    chain = LLMChain(
        transformers_model_path=model_size, 
        max_new_tokens=300, 
        api_key=api_key
    )
    return chain

# Generating the article based on user input
def generate_article(chain, prompt, temperature, top_p):
    response = chain.generate(prompt, temperature=temperature, top_p=top_p)
    return response

# Creating a Word document with the generated article and fetched image
def create_word_doc(title, content, image_url):
    doc = Document()
    doc.add_heading(title, 0)
    if image_url:
        image = Image.open(io.BytesIO(requests.get(image_url).content))
        image_stream = io.BytesIO()
        image.save(image_stream, format='JPEG')
        doc.add_picture(image_stream, width=Inches(6.0))
    doc.add_paragraph(content)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Main Streamlit app
def main():
    # Store LLM generated responses
    if "messages" not in st.session_state.keys():
        st.session_state.messages = [{"role": "assistant", "content": "How may I assist you today?"}]
    
    # Code for the rest of the Streamlit app follows here...
    # ...
