
from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests

def load_llm(max_tokens, prompt_template):
    llm = CTransformers(
        model="llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens=max_tokens,
        temperature=0.7
    )
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    return llm_chain

def get_src_original_url(query):
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': "iMn2jjJXgPCqmalZsrDxYA5WcLSyt1FgopsBxY4M8rUxRc4POC83rsR3",
    }
    params = {
        'query': query,
        'per_page': 1,
    }
    response = requests.get(url, headers=headers, params=params)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            src_original_url = photos[0]['src']['original']
            return src_original_url
        else:
            st.write("No photos found for the given query.")
    else:
        st.write(f"Error: {response.status_code}, {response.text}")
    return None

def save_image_from_url(url, filename="temp_image.jpg"):
    response = requests.get(url, stream=True)
    response.raise_for_status()
    with open(filename, 'wb') as file:
        for chunk in response.iter_content(8192):
            file.write(chunk)

def create_word_docx(user_input, paragraph, image_filename):
    doc = Document()
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    doc.add_heading('Image Input', level=1)
    doc.add_picture(image_filename, width=Inches(4))
    return doc

def main():
    st.title("Article Generator App using Llama 2")
    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")

    if len(user_input) > 0 and len(image_input) > 0:
        col1, col2, col3 = st.columns([1,2,1])
        with col1:
            st.subheader("Generated Content by Llama 2")
            prompt_template = "You are a digital marketing and SEO expert and your task is to write article so write an article on the given topic: {user_input}. The article must be under 800 words. The article should be be lengthy."
            llm_chain = load_llm(max_tokens=800, prompt_template=prompt_template)
            result = llm_chain(user_input)
            if len(result) > 0:
                st.info("Your article has been been generated successfully!")
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")

        with col2:
            st.subheader("Fetched Image")
            image_url = get_src_original_url(image_input)
            save_image_from_url(image_url)
            st.image(image_url)

        with col3:
            st.subheader("Final Article to Download")
            image_filename = "temp_image.jpg"
            doc = create_word_docx(user_input, result['text'], image_filename)
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

if __name__ == "__main__":
    main()
