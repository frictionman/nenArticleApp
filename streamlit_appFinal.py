from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests
import replicate

# Set up the Replicate API token
if 'REPLICATE_API_TOKEN' in st.secrets:
    replicate_api = st.secrets['REPLICATE_API_TOKEN']
else:
    st.error("Replicate API Token is missing in the Streamlit secrets!")
    raise Exception("Missing Replicate API Token")

os.environ['REPLICATE_API_TOKEN'] = replicate_api

# Loading the model
def load_llm(max_tokens, prompt_template):
    llm_model_identifier = 'a16z-infra/llama7b-v2-chat:4f0a4744c7295c024a1de15e1a63c880d3da035fa1f49bfd344fe076074c8eea'
    llm = CTransformers(
        model=llm_model_identifier,
        model_type="llama",
        max_new_tokens=max_tokens,
        temperature=0.7
    )
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    return llm_chain

def get_src_original_url(query):
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': st.secrets["PEXELS_API_KEY"],
    }
    params = {
        'query': query,
        'per_page': 1,
    }
    response = requests.get(url, headers=headers, params=params)
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            return photos[0]['src']['original']
    return None

def main():
    st.title("Article Generator App using Llama 7b")

    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")

    if user_input and image_input:
        with st.spinner("Generating your content..."):
            prompt_template = "You are a digital marketing and SEO expert and your task is to write an article on the topic: {user_input}. The article must be under 800 words."
            llm_chain = load_llm(max_tokens=800, prompt_template=prompt_template)
            result = llm_chain(user_input)

            image_url = get_src_original_url(image_input)

            col1, col2, col3 = st.columns([1, 2, 1])

            with col1:
                st.subheader("Generated Content")
                st.write(result)

            with col2:
                st.subheader("Fetched Image")
                if image_url:
                    st.image(image_url)
                else:
                    st.warning("Couldn't fetch an image for the provided topic.")

            with col3:
                st.subheader("Download Article")
                # Assuming `result` contains the generated text in a dictionary with key 'text'.
                generated_text = result.get('text', '')
                doc = Document()
                doc.add_heading('Generated Article', 0)
                doc.add_paragraph(generated_text)
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                st.download_button(
                    label='Download Article',
                    data=doc_buffer,
                    file_name='Generated_Article.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

if __name__ == "__main__":
    main()
