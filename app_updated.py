
# Developed by Nen AI

import requests

# Function to fetch photo from Pexels
def fetch_photo(query):
    # Placeholder for your Pexels API key
    api_key = 'YOUR_API_KEY'
    
    # Endpoint and headers for Pexels API
    url = 'https://api.pexels.com/v1/search'
    headers = {
        'Authorization': api_key,
    }
    params = {
        'query': query,
        'per_page': 1,
    }
    
    # Make the API request
    response = requests.get(url, headers=headers, params=params)
    
    # If successful, return the photo URL
    if response.status_code == 200:
        data = response.json()
        photos = data.get('photos', [])
        if photos:
            return photos[0]['src']['original']
    return None



# Developed by Nen AI


from langchain.llms import CTransformers
from langchain.chains import LLMChain
from langchain import PromptTemplate
import streamlit as st 
from docx import Document
from docx.shared import Inches
import io
from PIL import Image

# Function to load the Llama 2 model
def load_llm(max_tokens, prompt_template):
    # Initialize the Llama 2 model with given parameters
    llm = CTransformers(
        model = "llama-2-7b-chat.ggmlv3.q8_0.bin",
        model_type="llama",
        max_new_tokens = max_tokens,
        temperature = 0.7
    )
    
    # Set up a "chain" for processing prompts
    llm_chain = LLMChain(
        llm=llm,
        prompt=PromptTemplate.from_template(prompt_template)
    )
    return llm_chain

# Function to create a Word document with user input, generated article, and image
def create_word_docx(user_input, paragraph, image_input):
    doc = Document()
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    image_stream = io.BytesIO()
    image_input.save(image_stream, format='PNG')
    image_stream.seek(0)
    doc.add_picture(image_stream, width=Inches(4))
    return doc

# Streamlit UI setup

if "generation_count" not in st.session_state:
    st.session_state.generation_count = 0

if st.session_state.generation_count < 3:
    st.session_state.generation_count += 1
else:
    if 'REPLICATE_API_TOKEN' in st.secrets:
        st.success('API key already provided!', icon='✅')
        replicate_api = st.secrets['REPLICATE_API_TOKEN']
    else:
        replicate_api = st.text_input('Hey, you reached your limit for the day. Please enter your Replicate API token:', type='password')
def main():
    
# Store LLM generated responses
if "messages" not in st.session_state.keys():
    st.session_state.messages = [{"role": "assistant", "content": "How may I assist you today?"}]

# Display or clear chat messages
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.write(message["content"])

def clear_chat_history():
    st.session_state.messages = [{"role": "assistant", "content": "How may I assist you today?"}]
st.sidebar.button('Clear Chat History', on_click=clear_chat_history)
st.title("Article Generator using Llama 2 by Nen AI")
    
    # Take input from the user for article topic and image topic
    
selected_model = st.sidebar.selectbox('Choose a Llama model', ['Llama2-7B', 'Llama2-13B'], key='selected_model')
if selected_model == 'Llama2-7B':
    llm = 'Llama2-7B model path or identifier'
elif selected_model == 'Llama2-13B':
    llm = 'Llama2-13B model path or identifier'

temperature = st.sidebar.slider('temperature', min_value=0.01, max_value=5.0, value=0.7, step=0.01)
top_p = st.sidebar.slider('top_p', min_value=0.01, max_value=1.0, value=0.9, step=0.01)
user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")
    
    if user_input and image_input:
        # Display generated content and fetched image side by side
        col1, col2, col3 = st.columns([1,2,1])
        
        # Generate article using Llama 2
        with col1:
            prompt_template = f"
system_message = "You are an expert writing assistant specialized in writing blogs. Do not pretend to be 'User'. You only respond once as 'Assistant'."
 so write an article on the given topic: {user_input}. The article should be under 800 words."
            llm_chain = load_llm(max_tokens=800, prompt_template=prompt_template)
            result = llm_chain(user_input)
            if result:
                st.write(result)
            else:
                st.error("Your article couldn't be generated!")
        
        # Fetch image using the fetch_photo function
        with col2:
            image_url = fetch_photo(image_input)
            if image_url:
                st.image(image_url)
            else:
                st.write("No image found for the given topic.")
        
        # Allow user to download the final article with the image
        with col3:
            doc = create_word_docx(user_input, result['text'], Image.open(requests.get(image_url, stream=True).raw))
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            st.download_button(
                label='Download Word Document',
                data=doc_buffer,
                file_name='document.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )


st.write("Want to get the full functionality? Reach out to us and receive automated blogs daily in whatever language you want.")
if __name__ == "__main__":
    main()
