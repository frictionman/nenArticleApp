import streamlit as st
import os
from docx import Document
from docx.shared import Inches
import io
import requests
import replicate

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'generation_count' not in st.session_state:
    st.session_state.generation_count = 0

def generate_llama2_response(prompt_input, temperature, top_p, max_length):
    try:
        output = replicate.run('a16z-infra/llama13b-v2-chat:df7690f1994d94e96ad9d568eac121aecf50684a0b0963b25a41cc40061269e5',
                               input={'prompt': f'{prompt_input} Assistant: ',
                                      'temperature': temperature, 'top_p': top_p, 'max_length': max_length, 'repetition_penalty': 1})
        return output
    except Exception as e:
        st.error(f"An error occurred while generating the article: {e}")
        return None

def create_word_docx(user_input, paragraph):
    doc = Document()
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def main():
    st.title("Article Generator App using Llama 2")
    user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
    image_input = st.text_input("Please enter the topic for the image you want to fetch!")
    temperature = st.slider('Temperature', min_value=0.0, max_value=1.0, value=0.7)
    top_p = st.slider('Top P', min_value=0.0, max_value=1.0, value=0.9)
    max_length = st.slider('Max Length', min_value=50, max_value=500, value=500)

    if st.session_state.generation_count >= 5:
        st.warning('You have reached your free daily limit. Please upgrade for full functionality.')
        return

    if len(user_input) > 0 and len(image_input) > 0:
        col1, col2, col3 = st.columns([1,2,1])
        with col1:
            st.subheader("Generated Content by Llama 2")
            prompt_template = f"You are a digital marketing and SEO expert and your task is to write an article on the given topic: {user_input}. The article must be under 800 words."
            
            try:
                result = generate_llama2_response(prompt_template, temperature, top_p, max_length)
                
                if result:
                    st.info("Your article has been been generated successfully!")
                    st.write(result)
                    st.session_state.generation_count += 1
                    doc_buffer = create_word_docx(user_input, result)
                    st.download_button(
                        label='Download Word Document',
                        data=doc_buffer,
                        file_name='document.docx', 
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                else:
                    st.error("Your article couldn't be generated or is empty!")

            except Exception as e:
                st.error(f"An error occurred while generating the article: {e}")

if __name__ == "__main__":
    main()
